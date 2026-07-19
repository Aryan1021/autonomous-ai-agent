"""
Professional Microsoft Word document generator.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from agent.models import (
    DocumentContext,
    DocumentResult,
    DocumentSection,
    WorkflowContext,
)
from core.exceptions import (
    DocumentException,
    ValidationException,
)

from core.config import settings

from core.logging_config import get_logger

logger = get_logger(__name__)


class DocumentGenerator:
    """
    Generates a professional Microsoft Word document
    from a DocumentContext.
    """

    def __init__(self) -> None:
        """
        Initialize the document generator.
        """

        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)

    def generate(
        self,
        workflow: WorkflowContext,
        context: DocumentContext,
    ) -> DocumentResult:
        """
        Generate a professional DOCX document.

        Args:
            context:
                Structured document context.

        Returns:
            Metadata describing the generated document.

        Raises:
            ValidationException:
                If the supplied context is invalid.

            DocumentException:
                If document generation fails.
        """

        if not context.sections:
            raise ValidationException(
                "Document context contains no sections."
            )

        logger.info(
            "[%s] Generating Microsoft Word document.",
            workflow.request_id,
        )

        try:

            document = Document()

            self._configure_document(document)

            self._add_title(
                document=document,
                context=context,
            )

            self._add_goal(
                document=document,
                context=context,
            )

            self._add_assumptions(
                document=document,
                context=context,
            )

            self._add_sections(
                document=document,
                context=context,
            )

            document.add_page_break()

            if settings.SHOW_REFLECTION_REPORT:

                self._add_reflection(
                    document=document,
                    context=context,
                )

            self._add_workflow_summary(
                document=document,
                context=context,
            )

            self._add_footer(document)

            filename = self._create_filename(
                context.title,
            )

            path = self.output_dir / filename

            word_count = sum(
                len(section.content.split())
                for section in context.sections
            )

            document.save(path)

            logger.info(
                "[%s] Document successfully saved to '%s' (%d sections, %d words).",
                workflow.request_id,
                path,
                len(context.sections),
                word_count,
            )

            return DocumentResult(
                title=context.title,
                filename=filename,
                document_path=str(path),
                generated_at=datetime.now(),
                section_count=len(context.sections),
                word_count=word_count,
            )

        except ValidationException:
            raise

        except Exception as exc:

            logger.exception(
                "[%s] Failed to generate Microsoft Word document.",
                workflow.request_id,
            )

            raise DocumentException(
                "Document generation failed."
            ) from exc

    def _configure_document(
        self,
        document: Document,
    ) -> None:
        """
        Configure document-wide settings.
        """

        logger.info(
            "Configuring document."
        )

        core = document.core_properties

        core.author = "Autonomous AI Agent"
        core.title = "AI Generated Document"
        core.subject = "Autonomous AI Agent Output"
        core.comments = (
            "Generated automatically using the "
            "Autonomous AI Agent."
        )

        style = document.styles["Normal"]

        style.font.name = "Calibri"
        style.font.size = Pt(11)

    def _create_filename(
        self,
        title: str,
    ) -> str:
        """
        Generate a filesystem-safe filename.
        """

        filename = title.lower()

        filename = re.sub(
            r"[^a-z0-9]+",
            "_",
            filename,
        )

        filename = filename.strip("_")
        filename = filename[:60]

        return f"{filename}.docx"

    def _add_title(
        self,
        document: Document,
        context: DocumentContext,
    ) -> None:
        """
        Add the document title page.
        """

        logger.info(
            "Adding title page."
        )

        title = document.add_heading(
            context.title,
            level=0,
        )

        title.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        subtitle = document.add_paragraph()

        subtitle.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        subtitle.add_run(
            "Generated by Autonomous AI Agent"
        ).italic = True

        date = document.add_paragraph()

        date.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        date.add_run(
            datetime.now().strftime(
                "%d %B %Y %H:%M"
            )
        )

        document.add_page_break()

    def _add_goal(
        self,
        document: Document,
        context: DocumentContext,
    ) -> None:
        """
        Add the goal section.
        """

        logger.info(
            "Adding goal section."
        )

        heading = document.add_heading(
            "Goal",
            level=1,
        )

        heading.runs[0].font.size = Pt(16)

        paragraph = document.add_paragraph(
            context.goal,
        )

        paragraph.paragraph_format.space_after = Pt(12)

    def _add_assumptions(
        self,
        document: Document,
        context: DocumentContext,
    ) -> None:
        """
        Add planner assumptions.
        """

        logger.info(
            "Adding assumptions."
        )

        heading = document.add_heading(
            "Assumptions",
            level=1,
        )

        heading.runs[0].font.size = Pt(16)

        for assumption in context.assumptions:

            paragraph = document.add_paragraph(
                style="List Bullet",
            )

            paragraph.add_run(
                assumption,
            )

            paragraph.paragraph_format.space_after = Pt(6)

    def _add_sections(
        self,
        document: Document,
        context: DocumentContext,
    ) -> None:
        """
        Add all generated document sections.
        """

        logger.info(
            "Adding %d document sections.",
            len(context.sections),
        )

        for index, section in enumerate(
            context.sections,
            start=1,
        ):
            self._add_section(
                document=document,
                section=section,
                index=index,
            )

    def _add_section(
        self,
        document: Document,
        section: DocumentSection,
        index: int,
    ) -> None:
        """
        Add a single document section.
        """

        heading = document.add_heading(
            f"{index}. {section.heading}",
            level=1,
        )

        heading.runs[0].font.size = Pt(16)

        paragraph = document.add_paragraph(
            section.content,
        )

        paragraph.paragraph_format.space_after = Pt(12)

    def _add_reflection(
        self,
        document: Document,
        context: DocumentContext,
    ) -> None:
        """
        Add the reflection report.
        """

        logger.info(
            "Adding reflection report."
        )

        reflection = context.reflection

        # ---------------------------------------------------------
        # Heading
        # ---------------------------------------------------------

        heading = document.add_heading(
            "Reflection Report",
            level=1,
        )

        heading.runs[0].font.size = Pt(16)

        # ---------------------------------------------------------
        # Metrics
        # ---------------------------------------------------------

        paragraph = document.add_paragraph()
        paragraph.add_run(
            "Overall Score: "
        ).bold = True
        paragraph.add_run(
            f"{reflection.overall_score}/100"
        )

        paragraph = document.add_paragraph()
        paragraph.add_run(
            "Confidence: "
        ).bold = True
        paragraph.add_run(
            f"{reflection.confidence}%"
        )

        paragraph = document.add_paragraph()
        paragraph.add_run(
            "Needs Revision: "
        ).bold = True
        paragraph.add_run(
            "Yes"
            if reflection.needs_revision
            else "No"
        )

        # ---------------------------------------------------------
        # Strengths
        # ---------------------------------------------------------

        document.add_heading(
            "Strengths",
            level=2,
        )

        if reflection.strengths:

            for strength in reflection.strengths:

                document.add_paragraph(
                    strength,
                    style="List Bullet",
                )

        else:

            document.add_paragraph(
                "No strengths identified."
            )

        # ---------------------------------------------------------
        # Issues
        # ---------------------------------------------------------

        document.add_heading(
            "Issues",
            level=2,
        )

        if reflection.issues:

            for index, issue in enumerate(
                reflection.issues,
                start=1,
            ):

                document.add_heading(
                    f"Issue {index}",
                    level=3,
                )

                paragraph = document.add_paragraph()
                paragraph.add_run(
                    "Severity: "
                ).bold = True
                paragraph.add_run(
                    issue.severity.upper()
                )

                paragraph = document.add_paragraph()
                paragraph.add_run(
                    "Title: "
                ).bold = True
                paragraph.add_run(
                    issue.title
                )

                paragraph = document.add_paragraph()
                paragraph.add_run(
                    "Description: "
                ).bold = True
                paragraph.add_run(
                    issue.description
                )

                paragraph = document.add_paragraph()
                paragraph.add_run(
                    "Evidence: "
                ).bold = True
                paragraph.add_run(
                    issue.evidence
                )

                # Space before next issue
                document.add_paragraph()

        else:

            document.add_paragraph(
                "No issues identified."
            )

        # ---------------------------------------------------------
        # Suggestions
        # ---------------------------------------------------------

        document.add_heading(
            "Suggestions",
            level=2,
        )

        if reflection.suggestions:

            for suggestion in reflection.suggestions:

                paragraph = document.add_paragraph(
                    style="List Bullet",
                )

                paragraph.add_run(
                    suggestion.title
                ).bold = True

                document.add_paragraph(
                    suggestion.description
                )

                # Space before next suggestion
                document.add_paragraph()

        else:

            document.add_paragraph(
                "No suggestions."
            )

        # ---------------------------------------------------------
        # Reflection Summary
        # ---------------------------------------------------------

        document.add_heading(
            "Reflection Summary",
            level=2,
        )

        document.add_paragraph(
            reflection.summary,
        )

    def _add_workflow_summary(
        self,
        document: Document,
        context: DocumentContext,
    ) -> None:
        """
        Add workflow summary.
        """

        logger.info(
            "Adding workflow summary."
        )

        heading = document.add_heading(
            "Workflow Summary",
            level=1,
        )

        heading.runs[0].font.size = Pt(16)

        paragraph = document.add_paragraph(
            context.summary,
        )

        paragraph.paragraph_format.space_after = Pt(12)

    def _add_footer(
        self,
        document: Document,
    ) -> None:
        """
        Add a footer to the document.
        """

        logger.info(
            "Adding footer."
        )

        section = document.sections[0]

        footer = section.footer

        paragraph = footer.paragraphs[0]

        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        paragraph.text = (
            "Generated by Autonomous AI Agent\n"
            "Confidential"
        )